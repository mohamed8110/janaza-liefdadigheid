import streamlit as st
from docx import Document
from datetime import date
import os

st.set_page_config(page_title="Liefdadigheidsverklaring - Janaza", layout="centered")
st.title("❤️ Liefdadigheidsverklaring")

st.markdown("Vul onderstaande gegevens in om een verklaring te genereren.")

with st.form("verklaring_form"):
    naam_contactpersoon = st.text_input("Naam contactpersoon")
    naam_overledene = st.text_input("Naam overledene")
    telefoon = st.text_input("Telefoonnummer")
    email = st.text_input("E-mailadres")
    naam_contact = st.text_input("Naam voor ondertekening")
    datum_mandaat = st.date_input("Datum ondertekening", value=date.today())
    bestandsnaam = st.text_input("Bestandsnaam voor document", value="liefdadigheidsverklaring")

    submitted = st.form_submit_button("📄 Genereer verklaring")

if submitted:
    template_path = "Template_verklaring.docx"
    if not os.path.exists(template_path):
        st.error("❗ Templatebestand ontbreekt.")
    else:
        doc = Document(template_path)
        vervangingen = {
            "<<NAAM_CONTACTPERSOON>>": naam_contactpersoon,
            "<<NAAM_OVERLEDENE>>": naam_overledene,
            "<<TELEFOON>>": telefoon,
            "<<EMAIL>>": email,
            "<<DATUM_MANDAAT>>": datum_mandaat.strftime("%d/%m/%Y"),
            "<<NAAM_CONTACT>>": naam_contact
        }

        # Replace in normal paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                for key, val in vervangingen.items():
                    if key in run.text:
                        run.text = run.text.replace(key, val)

        # Replace in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for key, val in vervangingen.items():
                                if key in run.text:
                                    run.text = run.text.replace(key, val)

        output_docx = f"{bestandsnaam}.docx"
        doc.save(output_docx)

        with open(output_docx, "rb") as f:
            st.success("✅ Verklaring succesvol gegenereerd!")
            st.download_button("📥 Download .docx", f, file_name=output_docx, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
